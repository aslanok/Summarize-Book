import streamlit as st
import markdown2
import pdfplumber
import docx
from io import StringIO
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from chromadb import Client, Settings
from uuid import uuid4
from langchain_community.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import Chroma
import os

# ----------------------------
# Function Definitions
# ----------------------------

def read_pdf(file):
    """Read and extract text from a PDF file."""
    try:
        pdf = pdfplumber.open(file)
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        pdf.close()
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return None

def read_docx(file):
    """Read and extract text from a DOCX file."""
    try:
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return None

def read_txt(file):
    """Read and extract text from a TXT file."""
    try:
        stringio = StringIO(file.getvalue().decode("utf-8"))
        text = stringio.read()
        return text
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
        return None

def read_md(file):
    """Read and extract text from a Markdown (MD) file."""
    try:
        markdown_content = file.getvalue().decode("utf-8")
        text = markdown2.markdown(markdown_content)
        return text
    except Exception as e:
        st.error(f"Error reading MD file: {e}")
        return None

def process_file(file):
    """Process the uploaded file based on its extension."""
    if file.name.endswith(".pdf"):
        return read_pdf(file)
    elif file.name.endswith(".docx"):
        return read_docx(file)
    elif file.name.endswith(".txt"):
        return read_txt(file)
    elif file.name.endswith(".md"):
        return read_md(file)
    else:
        st.error("Unsupported file type!")
        return None

def split_text(documents: list):
    """Split documents into smaller chunks for processing."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=400,
        chunk_overlap=50,
        length_function=len,
        add_start_index=True,
    )
    chunks = text_splitter.split_documents(documents)
    st.write(f"Split {len(documents)} document(s) into {len(chunks)} chunk(s).")
    return chunks

# ----------------------------
# Streamlit App
# ----------------------------

# Sidebar for file upload and configuration
st.sidebar.title("📥 Upload and Configure")
uploaded_file = st.sidebar.file_uploader("Upload a document", type=["pdf", "docx", "txt", "md"])

# Main interface
st.title("📄 Ask Your File")
st.write("""
Welcome to the **Ask Your File** Application! 

**Features:**
- **Multiple File Formats**: Upload PDF, DOCX, TXT, or Markdown files.
- **Text Summarization**: Get automatic summaries of your documents.
- **Question-Answering**: Ask specific questions about your document's content.
- **Local Processing**: Ensures privacy and high performance by processing data locally using Ollama's LLM.
""")

if uploaded_file is not None:
    with st.spinner('📂 Processing the uploaded file...'):
        # Process the uploaded file
        text_content = process_file(uploaded_file)
        if text_content is None:
            st.error("Failed to process the uploaded file.")
            st.stop()

        # Display file details
        st.sidebar.success("File uploaded successfully!")
        st.write(f"**Uploaded File:** {uploaded_file.name}")
        st.write(f"**File Size:** {uploaded_file.size / 1024:.2f} KB")
        st.write("The file has been processed successfully.")

        # Wrap the text in a Document object
        documents = [Document(page_content=text_content, metadata={"source": uploaded_file.name})]

        # Generate a unique identifier for the document
        document_id = str(uuid4())

        # Split the text into chunks
        chunks = split_text(documents)

        # Load embedding model
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

        # Embed the text chunks
        chunk_embeddings = []
        for chunk in chunks:
            embedding = embedding_model.encode(chunk.page_content)
            chunk_embeddings.append({
                "embedding": embedding,
                "metadata": chunk.metadata,
                "content": chunk.page_content
            })

        # Initialize ChromaDB client and collection using Streamlit's session state
        if "chroma_client" not in st.session_state:
            # Create a new client instance and store it in session state
            chroma_dir = f'./chromadb_data/{document_id}'
            os.makedirs(chroma_dir, exist_ok=True)
            st.session_state.chroma_client = Client(Settings(persist_directory=chroma_dir))
            st.session_state.chroma_collection_name = f'doc_{document_id}'
            st.session_state.chroma_collection = st.session_state.chroma_client.create_collection(st.session_state.chroma_collection_name)
        else:
            chroma_dir = f'./chromadb_data/{document_id}'

        # Add embeddings and metadata to the collection stored in session state
        with st.spinner('📊 Indexing the document...'):
            for chunk_data in chunk_embeddings:
                doc_id = str(uuid4())
                embedding_list = chunk_data["embedding"].tolist()
                st.session_state.chroma_collection.add(
                    ids=[doc_id],
                    documents=[chunk_data["content"]],
                    embeddings=[embedding_list],
                    metadatas=[chunk_data["metadata"]]
                )

        # Initialize the LLM and QA chain
        llm = Ollama(model="gemma2:2b")
        prompt = """
        1. Use the following context to answer the question at the end.
        2. Be precise and avoid speculation. If the information isn't clear, say "I don't know."
        3. Provide a concise, 2-3 sentence answer.

        Context: {context}

        Question: {question}

        Accurate Answer:"""

        QA_CHAIN_PROMPT = PromptTemplate.from_template(prompt)

        def manual_llm_chain(llm, prompt_template, question):
            context = prompt_template.format(context=question["context"], question=question["question"])
            return llm(context)

        # Initialize embedding model for vector store
        embedding_store_model = HuggingFaceEmbeddings(model_name='all-MiniLM-L6-v2')

        # Create or load the vector store
        vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=embedding_store_model,
            persist_directory=chroma_dir,
        )

        # Initialize the Retrieval QA chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",  # Using the "stuff" method for simplicity
            retriever=vector_store.as_retriever(),
            verbose=True
        )

    # User input for query
    st.subheader("❓ Ask a Question or Request a Summary")
    query = st.text_input("Enter your query here:", "Can you provide a summary of the document?")

    if st.button("🔍 Generate Answer"):
        if query.strip() == "":
            st.warning("Please enter a valid query.")
        else:
            with st.spinner('🤖 Generating your answer...'):
                try:
                    # Retrieve relevant documents
                    retrieved_docs = qa_chain.retriever.get_relevant_documents(query)
                    if not retrieved_docs:
                        st.info("No relevant information found in the document.")
                    else:
                        # Format the query for the LLM
                        formatted_query = {
                            "context": " ".join([doc.page_content for doc in retrieved_docs]),
                            "question": query
                        }
                        # Generate the answer using the LLM
                        result = manual_llm_chain(llm, QA_CHAIN_PROMPT, formatted_query)
                        st.subheader("📝 Generated Answer:")
                        st.write(result)
                except Exception as e:
                    st.error(f"An error occurred while generating the answer: {e}")

    # Optional: Display retrieved documents (for debugging or transparency)
    with st.expander("📄 View Retrieved Document Sections"):
        if 'retrieved_docs' in locals() and retrieved_docs:
            for i, doc in enumerate(retrieved_docs, 1):
                st.markdown(f"**Section {i}:** {doc.page_content}")

else:
    st.info("🛠️ Please upload a document to get started.")

# Footer
st.markdown("---")
st.markdown("""
🔒 **Privacy:** All processing is done locally on your machine, ensuring your data remains private.

📂 **Support:** If you encounter any issues or have questions, please visit the [GitHub repository](https://github.com/aslanok/Summarize-Book).

🚀 **Powered by:** Streamlit, LangChain, ChromaDB, Ollama(Local LLM Model)
""")
